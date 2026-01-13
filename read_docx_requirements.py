
import docx
import os

docx_path = r'c:\Users\HP\OneDrive\Desktop\SMR_AGENT\Study Characteristics Table.docx'
output_path = r'c:\Users\HP\OneDrive\Desktop\SMR_AGENT\requirements_output.txt'

def read_docx(file_path):
    doc = docx.Document(file_path)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("--- Paragraphs ---\n")
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + "\n")
        
        f.write("\n--- Tables ---\n")
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                f.write(row_text + "\n")

if __name__ == "__main__":
    if os.path.exists(docx_path):
        read_docx(docx_path)
    else:
        print(f"File not found: {docx_path}")
